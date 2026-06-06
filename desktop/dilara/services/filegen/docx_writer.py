"""Word (.docx) dosyası üretimi."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from dilara.core.logging import logger


def write_docx(
    output_path: Path,
    title: str,
    body: str,
    *,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
) -> Path:
    """
    Markdown benzeri body'yi Word'e dönüştür.

    Desteklenen sözdizimi:
    * ``# Başlık``       -> H1
    * ``## Alt Başlık``  -> H2
    * ``- madde``        -> bullet list
    * ``1. madde``       -> numbered list
    * Boş satır          -> paragraf ayırıcı
    """
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        logger.error(f"python-docx yüklenemedi: {e}")
        raise

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    if title:
        doc.add_heading(title, level=0)
    if subtitle:
        doc.add_paragraph(subtitle).italic = True
    if author:
        doc.add_paragraph(f"Hazırlayan: {author}").italic = True

    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in (". ", ") "):
            doc.add_paragraph(stripped[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(stripped)

    doc.save(str(output_path))
    return output_path
